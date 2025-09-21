from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
import os
from .AI_Automation import Automation

class create_resume:
    
    def __init__(self,personal_detail,education_detail,skills_detail,projects_detail,internships_detail,certifications_detail):
        self.personal_detail = personal_detail
        self.education_detail = education_detail
        self.skills_detail = skills_detail
        self.projects_detail = projects_detail
        self.internships_detail = internships_detail
        # self.extracurricular_detail = extracurricular_detail
        self.certifications_detail = certifications_detail
        
        # -------------------------------------
        # Header details
        self.firstname = self.personal_detail.first_name
        self.lastname = self.personal_detail.last_name
        self.email = self.personal_detail.email
        self.phone  = self.personal_detail.phone
        self.linkedin = self.personal_detail.linkedin
        self.github = self.personal_detail.github
        
        
        # -------------------------------------
        # Education details
        self.clg_name = self.education_detail.college_name
        self.clg_dist = self.education_detail.college_district
        self.clg_passout = self.education_detail.college_passed_out_year
        self.degree = self.education_detail.degree
        self.branch = self.education_detail.branch
        self.cgpa = self.education_detail.cgpa
        self.cgpa_sem = self.education_detail.cgpa_semester
        self.hsc_name = self.education_detail.hsc_scl_name
        self.hsc_dist = self.education_detail.hsc_scl_district
        hsc_state = self.education_detail.hsc_scl_state
        self.hsc_poy = self.education_detail.hsc_passed_out_year
        self.hsc_percentage = self.education_detail.hsc_percentage
        self.sslc_name = self.education_detail.sslc_scl_name
        self.sslc_dist = self.education_detail.sslc_scl_district
        sslc_state = self.education_detail.sslc_scl_state
        self.sslc_percentage = self.education_detail.sslc_percentage
        self.sslc_poy = self.education_detail.sslc_passed_out_year
        
        
        # -------------------------------------
        # skills details
        self.prog_lang = self.skills_detail.prog_languages
        try:
            self.tool_tech  =self.skills_detail.Tools_technologies
        except Exception as e:
            self.tool_tech = None
            raise Exception(e)

        # print(f"tool_tech: {tool_tech}")
        
        
        # -------------------------------------
        # Projects details
        try:
            self.proj_count = self.projects_detail.project_count
            self.proj_names = self.projects_detail.project_name
            self.proj_keywords = self.projects_detail.project_keywords
            self.proj_descs = self.projects_detail.project_description

        except AttributeError as e:
            self.proj_count = 0
            print("No projects found",e)
        except Exception as e:
            print("Error occurred while fetching project details:", e)

        # print(f"proj_count: {proj_count}")
        # print(f"proj_names: {proj_names}")
        # print(f"proj_keywords: {proj_keywords}")
        # print(f"proj_descs: {proj_descs}")

        
        
        # -------------------------------------
        # Internship details
        try:
            self.comp_name = self.internships_detail.company_name
            self.role = self.internships_detail.role
            self.location = self.internships_detail.location
            self.duration = self.internships_detail.duration
            self.start_date = self.internships_detail.start_date
            self.end_date = self.internships_detail.end_date
        except AttributeError as e:
            self.comp_name = None
            print("No details available for internships:", e)
        except Exception as e:
            print("Error occurred while fetching internship details:", e)
            
        # -------------------------------------
        # Certifications details
        try:
            self.certificate_name  = self.certifications_detail.certification_name
        except AttributeError as e:
            self.certificate_name = None
            print("No details available for certifications:", e)
        except Exception as e:
            print("Error occurred while fetching certification details:", e)

            
        # Create document
        self.doc = Document()

    def alignment (self):

        # Set page margins
        for section in self.doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

    # Function to add hyperlinks
    def add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Blue underline style
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0000FF")
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), "single")
        rPr.append(u)

        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    # Helper to add section title
    def add_section_title(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(18)  # 18pt line spacing
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        

        # Add a bottom border (line under title)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')   # line style
        bottom.set(qn('w:sz'), '6')         # thickness
        bottom.set(qn('w:space'), '1')      # margin
        bottom.set(qn('w:color'), '000000') # black
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    # Helper to add bullet list
    def add_bullets(self,items):
        for item in items:
            p = self.doc.add_paragraph(item, style="List Bullet")
            p.style.font.name = "Times New Roman"
            p.style.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    # ---------- HEADER ----------
    def header(self):
        
        header = self.doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.add_run(f"{self.firstname} {self.lastname}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        p1 = self.doc.add_paragraph()
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r1 = p1.add_run("Karur, Tamil Nadu    ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        # self.add_hyperlink(p1, "karthikeyansa8@gmail.com", "mailto:karthikeyansa8@gmail.com")
        self.add_hyperlink(p1,f"{self.email}",f"mailto:{self.email}")

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # r2 = p2.add_run("7358996885    ")
        r2 = p2.add_run(f"{self.phone}    ")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        # self.add_hyperlink(p2, "LinkedIn", "https://www.linkedin.com/in/karthikeyan-s-a-4ab1a527b")
        self.add_hyperlink(p2, "LinkedIn", self.linkedin)
        p2.add_run("    ")
        # self.add_hyperlink(p2, "GitHub", "https://github.com/karthikeyansa8")
        self.add_hyperlink(p2, "GitHub", self.github)

    # ---------- SUMMARY ----------
    def summary(self):
        
        self.add_section_title("SUMMARY")
        
        input = f"Programming Languages: {self.prog_lang}, Tools & Technologies: {self.tool_tech}"
        Summary_text = Automation().summary(input)
        
        summary = self.doc.add_paragraph(
            # "Moderate-level experience in Python and Django, with the ability to build web applications and write efficient scripts. "
            # "Comfortable working with core programming concepts such as functions, loops, data types, and object-oriented programming. "
            # "I also have a basic understanding of Java, covering syntax and fundamental principles. "
            # "Always eager to learn and improve through practical, hands-on projects."
            Summary_text
        )
        summary.paragraph_format.line_spacing = Pt(18)  # 18pt line spacing

        summary.runs[0].font.name = "Times New Roman"
        summary.runs[0].font.size = Pt(11)

    # ---------- EDUCATION ----------
    def education(self):
        self.add_section_title("EDUCATION")

        # ---- Helper: heading with bold name, normal district, year inside margins ----
        def add_heading_with_year(name_text, district_text, year_text):
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(18)

            # Right tab stop set to the **usable text width** (inside margins)
            section = self.doc.sections[0]
            usable_width = section.page_width - section.left_margin - section.right_margin
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

            # Bold name
            run_name = p.add_run(name_text)
            run_name.bold = True
            run_name.font.name = "Times New Roman"
            run_name.font.size = Pt(11)

            # Normal district
            run_dist = p.add_run(f", {district_text}\t")
            run_dist.font.name = "Times New Roman"
            run_dist.font.size = Pt(11)

            # Year (right-aligned at tab stop)
            run_year = p.add_run(year_text)
            run_year.font.name = "Times New Roman"
            run_year.font.size = Pt(11)

        # ---- Helper: indented bullet ----
        def add_bullet(text):
            p = self.doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.left_indent = Inches(0.5)  # one tab indent
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

        # College
        add_heading_with_year(
            self.clg_name.upper(),
            self.clg_dist.title(),
            self.clg_passout
        )
        add_bullet(f"{self.degree.title()} in {self.branch.title()}")
        add_bullet(f"CGPA : {self.cgpa}/10 upto {self.cgpa_sem} semester")

        # HSC
        add_heading_with_year(
            self.hsc_name.upper(),
            self.hsc_dist.title(),
            self.hsc_poy
        )
        add_bullet(f"HSC – PERCENTAGE : {self.hsc_percentage}")

        # SSLC
        add_heading_with_year(
            self.sslc_name.upper(),
            self.sslc_dist.title(),
            self.sslc_poy
        )
        add_bullet(f"SSLC – PERCENTAGE : {self.sslc_percentage}")                    
                    

    # ---------- SKILLS ----------

    def skills(self):
        self.add_section_title("SKILLS")

        # Programming Languages
        p1 = self.doc.add_paragraph()
        p1.paragraph_format.line_spacing = Pt(18)
        run1_title = p1.add_run("Programming Languages: ")
        run1_title.bold = True
        run1_title.font.name = "Times New Roman"
        run1_title.font.size = Pt(11)

        run1_val = p1.add_run(self.prog_lang.title())
        run1_val.font.name = "Times New Roman"
        run1_val.font.size = Pt(11)

        # Tools and Technologies
        if self.tool_tech and self.tool_tech is not None:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.line_spacing = Pt(18)
            run2_title = p2.add_run("Tools And Technologies: ")
            run2_title.bold = True
            run2_title.font.name = "Times New Roman"
            run2_title.font.size = Pt(11)

            run2_val = p2.add_run(self.tool_tech.title())
            run2_val.font.name = "Times New Roman"
            run2_val.font.size = Pt(11)


    # ---------- PROJECTS ----------
    def projects(self):
        
        
        if self.proj_count > 0:
            projects = []
            for i in range(1, self.proj_count + 1):
                name = self.proj_names.get(f"project{i}_name", "")
                keywords = self.proj_keywords.get(f"project{i}_keywords", "")
                desc = self.proj_descs.get(f"project{i}_description", "")
        
                title = f"{name.title()} | {keywords.title()}"
                                
                projects.append((title, desc))

            
            self.add_section_title("PROJECTS")
            
            for title, tasks in projects:
                t = self.doc.add_paragraph()
                t.paragraph_format.line_spacing = Pt(18)  # 18pt line spacing
                tr = t.add_run(title)
                tr.bold = True
                tr.font.name = "Times New Roman"
                tr.font.size = Pt(11)
                
                # Project Tasks (single paragraph, not bullets)
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = Pt(18)
                pr = p.add_run(tasks.strip())
                pr.font.name = "Times New Roman"
                pr.font.size = Pt(11)

    # ---------- INTERNSHIP ----------
  

    def internship(self):
        if self.comp_name:
            self.add_section_title("INTERNSHIP")

            # Main heading line
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(18)

            # Company name – bold
            run_company = p.add_run(f"{self.comp_name.title()} ")
            run_company.bold = True
            run_company.font.name = "Times New Roman"
            run_company.font.size = Pt(11)

            # Location – normal
            run_loc = p.add_run(f"({self.location.title()})  ")
            run_loc.font.name = "Times New Roman"
            run_loc.font.size = Pt(11)

            # Role – bold
            run_role = p.add_run(f"Role: {self.role.title()} ")
            run_role.bold = True
            run_role.font.name = "Times New Roman"
            run_role.font.size = Pt(11)

            # Duration – normal, inside page margins with a separator
            run_period = p.add_run(f"| {self.start_date} - {self.end_date}")
            run_period.font.name = "Times New Roman"
            run_period.font.size = Pt(11)

            # Bullet points describing the internship
            input_text = f"Worked as a {self.role} intern at {self.comp_name}"
            intern_desc = Automation().intern_description(input_text)
            intern_desc = intern_desc.split("||")

            for desc in intern_desc:
                bullet = self.doc.add_paragraph(desc, style='List Bullet')
                bullet.paragraph_format.line_spacing = Pt(18)
                bullet.paragraph_format.left_indent = Inches(0.5)  # indent bullets
                for r in bullet.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

    # # ---------- LEADERSHIP ----------
    # def leadership(self):
    #         self.add_section_title("LEADERSHIP / EXTRACURRICULAR")
    #         self.doc.add_paragraph("Logo Design Contest Participant\nKS Rangasamy College of Technology | 2024 | National Level").runs[0].font.name = "Times New Roman"
    #         self.add_bullets([
    #             "Contributed creative and innovative designs to the Service Motto Volunteers contest.",
    #             "Recognized for dedication and excellence in graphic design."
    #         ])

    # ---------- CERTIFICATIONS ----------
    def certifications(self):      
        
        
        if self.certificate_name is not None:
            self.add_section_title("CERTIFICATIONS")

            # self.add_bullets(certificate_name)
            for item in self.certificate_name:
                p = self.doc.add_paragraph(item.title(), style="List Bullet")
                p.paragraph_format.line_spacing = Pt(18)  # 18pt line spacing
                p.style.font.name = "Times New Roman"
                p.style.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # Save file
    def save(self):
        os.makedirs("Resumes", exist_ok=True)
        self.doc.save(f"Resumes/{self.personal_detail.first_name}_{self.personal_detail.last_name}_Resume.docx")
        print("Resume created successfully!")
        # self.stdout.write(self.style.SUCCESS("Resume created successfully!"))

    
    def main(self):
        self.alignment()
        self.header()
        self.summary()
        self.education()
        self.skills()
        self.projects()
        self.internship()
        # self.leadership()
        self.certifications()
        self.save()